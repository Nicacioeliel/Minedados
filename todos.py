# ESTE CÓDIGO ELE EXTRAI DADOS DE INCORPORAÇÃO IMOBILIÁRIA DE APARTAMENTOS (BLOCO E TORRE) E DE CASA
# PARA APARTAMENTOS DE TORRE, ESSE É O PADRÃO: Apartamento 001, tipo A, da Torre 01, 
# PARA APARTAMENTOS DE BLOCO, ESSE É O PADRÃO: APARTAMENTO 001 – BLOCO 01

import streamlit as st
import os
import ocrmypdf
import pandas as pd
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document

# Configuração do Tesseract (ajuste se necessário)
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
os.environ['TESSDATA_PREFIX'] = r'C:\Program Files\Tesseract-OCR\tessdata'

# Funções de extração (adaptadas do extrator_unificado.py)
def extrair_texto_pdf(file):
    reader = PdfReader(file)
    texto = ""
    for page in reader.pages:
        texto += page.extract_text() + "\n"
    return texto

def identificar_tipo_documento(texto):
    import re
    if re.search(r"Apartamento\\s+\\d+,\\s*tipo\\s+[A-Z],\\s+da\\s+Torre", texto, re.IGNORECASE):
        return "torre"
    elif re.search(r"Apartamento\\s+\\d+,\\s*TIPO\\s+[A-Z],\\s*do\\s+Bloco\\s+\\d+", texto, re.IGNORECASE):
        return "bloco"
    elif re.search(r"Apartamento\\s+\\d+\\s*[-–]\\s*Bloco\\s+\\d+", texto, re.IGNORECASE):
        return "bloco"
    elif re.search(r"(?:CASA|Casa)[ \n]?[Nn]?[°º]?[ ]?\\d{2}", texto):
        return "casa"
    return "desconhecido"

# (Inclua aqui as funções extrair_torre, extrair_bloco, extrair_casas, salvar_excel do extrator_unificado.py)
# ...
# Por brevidade, não repito todas aqui, mas elas devem ser copiadas do seu extrator_unificado.py

def extrair_torre(doc):
    dados = []
    import re
    regex_torre = re.compile(
        r"Apartamento\s+(\d+),\s+tipo\s+([A-Z]),\s+da\s+Torre\s+(\d+).*?"
        r"área privativa principal de ([\d,]+)m².*?"
        r"(?:área privativa acessória de [\d,]+m².*?)?"
        r"área privativa total de ([\d,]+)m².*?"
        r"área de uso comum de ([\d,]+)m².*?"
        r"área real total de ([\d,]+)m².*?"
        r"fração ideal.*?de ([\d,]+).*?ou ([\d,]+)m²",
        re.DOTALL | re.IGNORECASE
    )
    for p in doc.paragraphs:
        texto = p.text.strip()
        match = regex_torre.search(texto)
        if match:
            numero, tipo, torre, privativa, total, comum, real, fracao, terreno = match.groups()
            descricao = ""
            idx_localizado = texto.lower().find('localizado')
            idx_ultimo_ponto = texto.rfind('.')
            if idx_localizado != -1 and idx_ultimo_ponto != -1 and idx_ultimo_ponto > idx_localizado:
                descricao = texto[idx_localizado:idx_ultimo_ponto+1].replace('\n', ' ').strip()
            dados.append({
                "Formato": "Torre",
                "Apartamento": numero,
                "Tipo": tipo,
                "Torre/Bloco": torre,
                "Área Privativa (m²)": privativa.replace(",", "."),
                "Área Comum (m²)": comum.replace(",", "."),
                "Área Total (m²)": real.replace(",", "."),
                "Fração Ideal (%)": fracao.replace(",", "."),
                "Área Terreno (m²)": terreno.replace(",", "."),
                "Descrição": descricao
            })
    return pd.DataFrame(dados)

def extrair_bloco(doc):
    dados = []
    import re
    regex_bloco2 = re.compile(
        r"Apartamento\s+(\d+),\s*TIPO\s*([A-Z]),\s*do\s+Bloco\s+(\d+),.*?Áreas:.*?"
        r"área privativa principal de ([\d,.]+)m².*?"
        r"área privativa total de ([\d,.]+)m².*?"
        r"área de uso comum de ([\d,.]+)m².*?"
        r"área real total de ([\d,.]+)m².*?"
        r"fração ideal.*?de ([\d,.]+).*?ou ([\d,.]+)m²",
        re.DOTALL | re.IGNORECASE
    )
    regex_bloco1 = re.compile(
        r"APARTAMENTO\s+(\d+)\s*[-–]\s*BLOCO\s+(\d+):.*?"
        r"áreas:\s*privativa real de ([\d,]+)m²,\s*"
        r"área de uso comum real de ([\d,]+)m²,\s*"
        r"perfazendo uma área total real de ([\d,]+)m².*?"
        r"área equivalente de construção igual a ([\d,]+)m².*?"
        r"fração ideal.*?([0-9,.]+)%",
        re.DOTALL | re.IGNORECASE
    )
    for p in doc.paragraphs:
        texto = p.text.strip()
        match2 = regex_bloco2.search(texto)
        if match2:
            numero, tipo, bloco, privativa, total, comum, real, fracao, terreno = match2.groups()
            idx_localizado = texto.lower().find('localizado')
            descricao = ""
            if idx_localizado != -1:
                trecho = texto[idx_localizado:]
                idx_ultimo_ponto = trecho.rfind('.')
                if idx_ultimo_ponto != -1:
                    descricao = trecho[:idx_ultimo_ponto+1].replace('\n', ' ').strip()
                else:
                    descricao = trecho.strip()
            dados.append({
                "Formato": "Bloco",
                "Apartamento": numero,
                "Tipo": tipo,
                "Torre/Bloco": bloco,
                "Área Privativa (m²)": privativa.replace(",", "."),
                "Área Privativa Total (m²)": total.replace(",", "."),
                "Área Comum (m²)": comum.replace(",", "."),
                "Área Total (m²)": real.replace(",", "."),
                "Fração Ideal (%)": fracao.replace(",", "."),
                "Área Terreno (m²)": terreno.replace(",", "."),
                "Descrição": descricao
            })
            continue
        match1 = regex_bloco1.search(texto)
        if match1:
            numero, bloco, privativa, comum, total, equivalente, fracao = match1.groups()
            idx_localizado = texto.lower().find('localizado')
            descricao = ""
            if idx_localizado != -1:
                trecho = texto[idx_localizado:]
                idx_ultimo_ponto = trecho.rfind('.')
                if idx_ultimo_ponto != -1:
                    descricao = trecho[:idx_ultimo_ponto+1].replace('\n', ' ').strip()
                else:
                    descricao = trecho.strip()
            dados.append({
                "Formato": "Bloco",
                "Apartamento": numero,
                "Tipo": "",
                "Torre/Bloco": bloco,
                "Área Privativa (m²)": privativa.replace(",", "."),
                "Área Comum (m²)": comum.replace(",", "."),
                "Área Total (m²)": total.replace(",", "."),
                "Fração Ideal (%)": fracao.replace(",", "."),
                "Área Terreno (m²)": equivalente.replace(",", "."),
                "Descrição": descricao
            })
    return pd.DataFrame(dados)

def extrair_casas(doc):
    import re
    texto = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    blocos = re.split(r"(?:^|\n)(?:CASA|Casa)[ \n]?[Nn]?[°º]?[ ]?(\d{2})", texto)
    casas = [(blocos[i], blocos[i + 1]) for i in range(1, len(blocos), 2)]
    regex_area_terreno = re.compile(r"configuração.*?área total de *(\d+,\d+)")
    regex_area_construida = re.compile(r"área total construída da casa de *(\d+,\d+)")
    regex_area_comum = re.compile(r"área de uso comum real de *(\d+,\d+)")
    regex_area_total_real = re.compile(r"área total real de *(\d+,\d+)")
    regex_fracao_ideal = re.compile(r"fração ideal do terreno correspondente a *(\d+,\d+)\s?%")
    regex_descricao = re.compile(r"(Pavimento térreo:.*?)(?:\.\s*$|\Z)", re.DOTALL)
    dados = []
    for numero, conteudo in casas:
        area_terreno = re.search(regex_area_terreno, conteudo)
        area_construida = re.search(regex_area_construida, conteudo)
        area_comum = re.search(regex_area_comum, conteudo)
        area_total = re.search(regex_area_total_real, conteudo)
        fracao = re.search(regex_fracao_ideal, conteudo)
        descricao = re.search(regex_descricao, conteudo)
        if not descricao:
            idx_frente = conteudo.lower().find('frente')
            idx_ultimo_ponto = conteudo.rfind('.')
            if idx_frente != -1 and idx_ultimo_ponto != -1 and idx_ultimo_ponto > idx_frente:
                descricao = conteudo[idx_frente:idx_ultimo_ponto+1].replace('\n', ' ').strip()
            else:
                descricao = ""
        else:
            descricao = descricao.group(1).strip()
        dados.append({
            "Número da Casa": numero,
            "Área do Terreno (m²)": area_terreno.group(1) if area_terreno else "",
            "Área Construída (m²)": area_construida.group(1) if area_construida else "",
            "Área Comum Real (m²)": area_comum.group(1) if area_comum else "",
            "Área Total Real (m²)": area_total.group(1) if area_total else "",
            "Fração Ideal": (fracao.group(1) + " %") if fracao else "",
            "Descrição": descricao
        })
    return pd.DataFrame(dados)

def salvar_excel(df, filename):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    for col, header in enumerate(df.columns, 1):
        ws.cell(row=1, column=col, value=header)
    for row_idx, row in enumerate(df.values, 2):
        for col_idx, value in enumerate(row, 1):
            ws.cell(row=row_idx, column=col_idx, value=value)
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer

# Função para OCR de um arquivo PDF (upload direto)
def ocr_pdf_upload(uploaded_pdf):
    temp_input = "temp_input.pdf"
    temp_output = "temp_output.pdf"
    with open(temp_input, "wb") as f:
        f.write(uploaded_pdf.read())
    try:
        ocrmypdf.ocr(
            temp_input,
            temp_output,
            language="por",
            force_ocr=True,
            deskew=True
        )
        with open(temp_output, "rb") as f:
            return f.read()
    except Exception as e:
        st.error(f"Erro ao processar PDF: {e}")
        return None
    finally:
        if os.path.exists(temp_input):
            os.remove(temp_input)
        if os.path.exists(temp_output):
            os.remove(temp_output)

# Função para OCR em lote (pasta inteiroteor)
def ocr_pasta_inteira(entrada_dir, saida_dir):
    os.makedirs(saida_dir, exist_ok=True)
    resultados = []
    for arquivo in os.listdir(entrada_dir):
        if arquivo.lower().endswith(".pdf"):
            caminho_entrada = os.path.join(entrada_dir, arquivo)
            caminho_saida = os.path.join(saida_dir, arquivo)
            try:
                ocrmypdf.ocr(
                    caminho_entrada,
                    caminho_saida,
                    language="por",
                    force_ocr=True,
                    deskew=True
                )
                resultados.append((arquivo, caminho_saida, None))
            except Exception as e:
                resultados.append((arquivo, None, str(e)))
    return resultados

# --- INTERFACE STREAMLIT ---
st.set_page_config(page_title="Sistema Unificado: Incorporação & OCR de PDFs", layout="wide")
st.title("Sistema Unificado: Incorporação & OCR de PDFs")

opcao = st.sidebar.selectbox(
    "Escolha a funcionalidade:",
    ["Extração de Dados de Memorial", "OCR de PDF (Upload Direto)", "OCR de PDF (Pasta Inteira)"]
)

if opcao == "Extração de Dados de Memorial":
    st.header("Extração de Dados de Memorial de Incorporação")
    uploaded_file = st.file_uploader("Envie um arquivo .docx ou .pdf", type=["docx", "pdf"])
    if uploaded_file:
        if uploaded_file.name.lower().endswith(".docx"):
            doc = Document(uploaded_file)
            texto_completo = "\n".join(p.text for p in doc.paragraphs)
        elif uploaded_file.name.lower().endswith(".pdf"):
            texto_completo = extrair_texto_pdf(uploaded_file)
            class FakeDoc:
                def __init__(self, texto):
                    self.paragraphs = [type('p', (), {'text': t}) for t in texto.split('\n') if t.strip()]
            doc = FakeDoc(texto_completo)
        else:
            st.warning("Tipo de arquivo não suportado. Use .docx ou .pdf")
            st.stop()
        tipo = identificar_tipo_documento(texto_completo)
        if tipo == "desconhecido":
            st.warning("Tipo de documento não reconhecido. Verifique se o arquivo contém dados válidos.")
            st.stop()
        if tipo == "torre":
            df = extrair_torre(doc)
        elif tipo == "bloco":
            df = extrair_bloco(doc)
        elif tipo == "casa":
            df = extrair_casas(doc)
        if df.empty:
            st.warning("Nenhum dado foi extraído do documento. Verifique se o formato está correto.")
            st.stop()
        st.success(f"Documento identificado como: {tipo.upper()}")
        st.write(f"Total de registros encontrados: {len(df)}")
        st.dataframe(df)
        if not df.empty:
            buffer = salvar_excel(df, f"dados_{tipo}.xlsx")
            st.download_button(
                label="📥 Baixar Excel",
                data=buffer,
                file_name=f"dados_{tipo}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
elif opcao == "OCR de PDF (Upload Direto)":
    st.header("Transformar PDF Digitalizado em Pesquisável (OCR)")
    uploaded_pdf = st.file_uploader("Envie um PDF digitalizado", type=["pdf"])
    if uploaded_pdf:
        resultado = ocr_pdf_upload(uploaded_pdf)
        if resultado:
            st.success("PDF processado com sucesso! Baixe abaixo:")
            st.download_button(
                label="📥 Baixar PDF pesquisável",
                data=resultado,
                file_name="pdf_pesquisavel.pdf",
                mime="application/pdf"
            )
elif opcao == "OCR de PDF (Pasta Inteira)":
    st.header("OCR em Lote: Pasta 'inteiroteor'")
    entrada_dir = "inteiroteor"
    saida_dir = "pdfs_editaveis"
    if st.button("Processar todos os PDFs da pasta"):
        resultados = ocr_pasta_inteira(entrada_dir, saida_dir)
        for arquivo, caminho_saida, erro in resultados:
            if erro:
                st.error(f"Erro ao processar {arquivo}: {erro}")
            else:
                with open(caminho_saida, "rb") as f:
                    st.download_button(
                        label=f"Baixar {arquivo}",
                        data=f.read(),
                        file_name=arquivo,
                        mime="application/pdf"
                    ) 
