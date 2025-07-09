# ESTE CÓDIGO ELE EXTRAI DADOS DE INCORPORAÇÃO IMOBILIÁRIA DE APARTAMENTOS (BLOCO E TORRE) E DE CASA
# PARA APARTAMENTOS DE TORRE, ESSE É O PADRÃO: Apartamento 001, tipo A, da Torre 01, 
# PARA APARTAMENTOS DE BLOCO, ESSE É O PADRÃO: APARTAMENTO 001 – BLOCO 01

import re
import pandas as pd
import streamlit as st
from docx import Document
from io import BytesIO

st.title("Extrator de Dados de Memorial de Incorporação")

uploaded_file = st.file_uploader("Envie um arquivo .docx", type="docx")

def identificar_tipo_documento(texto):
    if re.search(r"Apartamento\s+\d+,\s+tipo\s+[A-Z],\s+da\s+Torre", texto, re.IGNORECASE):
        return "torre"
    elif re.search(r"Apartamento\s+\d+\s*[-–]\s*Bloco\s+\d+", texto, re.IGNORECASE):
        return "bloco"
    elif re.search(r"(?:CASA|Casa)[ \n]?[Nn]?[°º]?[ ]?\d{2}", texto):
        return "casa"
    return "desconhecido"

def extrair_torre(doc):
    dados = []
    texto_completo = "\n".join([p.text for p in doc.paragraphs])

    regex_torre = re.compile(
        r"Apartamento\s+(\d+),\s+tipo\s+([A-Z]),\s+da\s+Torre\s+(\d+).*?"
        r"área privativa principal de ([\d,]+)m².*?"
        r"(?:área privativa acessória de [\d,]+m².*?)?"
        r"área privativa total de ([\d,]+)m².*?"
        r"área de uso comum de ([\d,]+)m².*?"
        r"área real total de ([\d,]+)m².*?"
        r"fração ideal.*?de ([\d,]+).*?ou ([\d,]+)m²",
        re.DOTALL | re.IGNORECASE
    )

    for p in doc.paragraphs:
        texto = p.text.strip()
        match = regex_torre.search(texto)
        if match:
            numero, tipo, torre, privativa, total, comum, real, fracao, terreno = match.groups()
            # Encontrar a descrição de 'localizado' até o último ponto final
            idx_localizado = texto.lower().find('localizado')
            idx_ultimo_ponto = texto.rfind('.')
            descricao = ""
            if idx_localizado != -1 and idx_ultimo_ponto != -1 and idx_ultimo_ponto > idx_localizado:
                descricao = texto[idx_localizado:idx_ultimo_ponto+1].replace('\n', ' ').strip()
            dados.append({
                "Formato": "Torre",
                "Apartamento": numero,
                "Tipo": tipo,
                "Torre/Bloco": torre,
                "Área Privativa (m²)": privativa.replace(",", "."),
                "Área Comum (m²)": comum.replace(",", "."),
                "Área Total (m²)": real.replace(",", "."),
                "Fração Ideal (%)": fracao.replace(",", "."),
                "Área Terreno (m²)": terreno.replace(",", "."),
                "Descrição": descricao
            })
    return pd.DataFrame(dados)

def extrair_bloco(doc):
    dados = []
    # Padrão 1: APARTAMENTO 001 – BLOCO 01: Localizado ...
    regex_bloco1 = re.compile(
        r"APARTAMENTO\s+(\d+)\s*[-–]\s*BLOCO\s+(\d+):?", re.IGNORECASE
    )
    # Padrão 2: Apartamento 001, TIPO A, do Bloco 01, localizado ...
    regex_bloco2 = re.compile(
        r"Apartamento\s+(\d+),\s*TIPO\s*([A-Z]),\s*do\s*Bloco\s+(\d+),", re.IGNORECASE
    )
    for p in doc.paragraphs:
        texto = p.text.strip()
        # Tenta padrão 2 primeiro (mais específico)
        match2 = regex_bloco2.search(texto)
        if match2:
            numero, tipo, bloco = match2.groups()
            # Descrição: de 'localizado' ou 'assim descrito:' até o último ponto final
            idx_localizado = texto.lower().find('localizado')
            idx_assim_descrito = texto.lower().find('assim descrito:')
            idx_inicio = idx_localizado if idx_localizado != -1 else idx_assim_descrito
            idx_ultimo_ponto = texto.rfind('.')
            descricao = ""
            if idx_inicio != -1 and idx_ultimo_ponto != -1 and idx_ultimo_ponto > idx_inicio:
                descricao = texto[idx_inicio:idx_ultimo_ponto+1].replace('\n', ' ').strip()
            # Extrair áreas (padrão novo)
            regex_areas2 = re.compile(
                r"Áreas:\s*área privativa principal de ([\d,.]+)m².*?"
                r"área privativa total de ([\d,.]+)m².*?"
                r"área de uso comum de ([\d,.]+)m².*?"
                r"área real total de ([\d,.]+)m².*?"
                r"fração ideal de solo de ([\d,.]+).*?ou ([\d,.]+)m²",
                re.DOTALL | re.IGNORECASE
            )
            match_areas2 = regex_areas2.search(texto)
            if match_areas2:
                privativa, total, comum, real, fracao, terreno = match_areas2.groups()
                dados.append({
                    "Formato": "Bloco",
                    "Apartamento": numero,
                    "Tipo": tipo,
                    "Torre/Bloco": bloco,
                    "Área Privativa (m²)": privativa.replace(",", "."),
                    "Área Comum (m²)": comum.replace(",", "."),
                    "Área Total (m²)": real.replace(",", "."),
                    "Fração Ideal (%)": fracao.replace(",", "."),
                    "Área Terreno (m²)": terreno.replace(",", "."),
                    "Descrição": descricao
                })
            continue
        # Tenta padrão 1
        match1 = regex_bloco1.search(texto)
        if match1:
            numero, bloco = match1.groups()
            tipo = ""
            # Descrição: de 'localizado' até o último ponto final
            idx_localizado = texto.lower().find('localizado')
            idx_ultimo_ponto = texto.rfind('.')
            descricao = ""
            if idx_localizado != -1 and idx_ultimo_ponto != -1 and idx_ultimo_ponto > idx_localizado:
                descricao = texto[idx_localizado:idx_ultimo_ponto+1].replace('\n', ' ').strip()
            # Extrair áreas (padrão antigo)
            regex_areas1 = re.compile(
                r"áreas:\s*privativa real de ([\d,]+)m²,\s*"
                r"área de uso comum real de ([\d,]+)m²,\s*"
                r"perfazendo uma área total real de ([\d,]+)m².*?"
                r"área equivalente de construção igual a ([\d,]+)m².*?"
                r"fração ideal.*?([0-9,.]+)%",
                re.DOTALL | re.IGNORECASE
            )
            match_areas1 = regex_areas1.search(texto)
            if match_areas1:
                privativa, comum, total, equivalente, fracao = match_areas1.groups()
                dados.append({
                    "Formato": "Bloco",
                    "Apartamento": numero,
                    "Tipo": tipo,
                    "Torre/Bloco": bloco,
                    "Área Privativa (m²)": privativa.replace(",", "."),
                    "Área Comum (m²)": comum.replace(",", "."),
                    "Área Total (m²)": total.replace(",", "."),
                    "Fração Ideal (%)": fracao.replace(",", "."),
                    "Área Terreno (m²)": equivalente.replace(",", "."),
                    "Descrição": descricao
                })
    return pd.DataFrame(dados)

def extrair_casas(doc):
    texto = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    blocos = re.split(r"(?:^|\n)(?:CASA|Casa)[ \n]?[Nn]?[°º]?[ ]?(\d{2})", texto)
    casas = [(blocos[i], blocos[i + 1]) for i in range(1, len(blocos), 2)]

    regex_area_terreno = re.compile(r"configuração.*?área total de *(\d+,\d+)")
    regex_area_construida = re.compile(r"área total construída da casa de *(\d+,\d+)")
    regex_area_comum = re.compile(r"área de uso comum real de *(\d+,\d+)")
    regex_area_total_real = re.compile(r"área total real de *(\d+,\d+)")
    regex_fracao_ideal = re.compile(r"fração ideal do terreno correspondente a *(\d+,\d+)\s?%")

    dados = []
    for numero, conteudo in casas:
        area_terreno = re.search(regex_area_terreno, conteudo)
        area_construida = re.search(regex_area_construida, conteudo)
        area_comum = re.search(regex_area_comum, conteudo)
        area_total = re.search(regex_area_total_real, conteudo)
        fracao = re.search(regex_fracao_ideal, conteudo)
        # Descrição: de 'frente' até o último ponto final
        idx_frente = conteudo.lower().find('frente')
        idx_ultimo_ponto = conteudo.rfind('.')
        descricao = ""
        if idx_frente != -1 and idx_ultimo_ponto != -1 and idx_ultimo_ponto > idx_frente:
            descricao = conteudo[idx_frente:idx_ultimo_ponto+1].replace('\n', ' ').strip()
        dados.append({
            "Número da Casa": numero,
            "Área do Terreno (m²)": area_terreno.group(1) if area_terreno else "",
            "Área Construída (m²)": area_construida.group(1) if area_construida else "",
            "Área Comum Real (m²)": area_comum.group(1) if area_comum else "",
            "Área Total Real (m²)": area_total.group(1) if area_total else "",
            "Fração Ideal": (fracao.group(1) + " %") if fracao else "",
            "Descrição": descricao
        })

    return pd.DataFrame(dados)

if uploaded_file:
    doc = Document(uploaded_file)
    texto_completo = "\n".join(p.text for p in doc.paragraphs)
    tipo = identificar_tipo_documento(texto_completo)

    if tipo == "torre":
        df = extrair_torre(doc)
    elif tipo == "bloco":
        df = extrair_bloco(doc)
    elif tipo == "casa":
        df = extrair_casas(doc)
    else:
        st.warning("Tipo de documento não reconhecido.")
        st.stop()

    st.success(f"Documento identificado como: {tipo.upper()}")
    st.dataframe(df)

    buffer = BytesIO()
    df.to_excel(buffer, index=False)
    buffer.seek(0)

    st.download_button(
        label="📥 Baixar Excel",
        data=buffer,
        file_name=f"dados_{tipo}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
